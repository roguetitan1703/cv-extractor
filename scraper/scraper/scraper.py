import re, os
from os import path, listdir
from PyPDF2 import PdfReader
from docx import Document
import docx
from openpyxl import Workbook
from ultra_logger import Logger
from spire.doc import Document as SpireDoc
import zipfile

class Scraper:
    def __init__(self,log_file="scraper.log"):
        self.log = Logger("Scraper.py", f"{log_file}", True, True)
        self.EMAIL_REGEX = r'[^.]([a-zA-Z0-9._%+-]+[^.])@([a-zA-Z0-9]+(?:\.[a-zA-Z0-9]+)*)(\s*\.[a-zA-Z]{2,})'
        self.PHONE_REGEX = r"(?:\+?\d{1,3}[\s-]?)?(?:\(\d{3}\)|\d{3})[\s-]?\d{2,4}[\s-]?\d{2,4}"


    def extract_phone_numbers(self, text):
        """
        Extracts phone numbers from text, handling optional country code (+XX-).

        Args:
            text (str): The text to search for phone numbers.

        Returns:
            list: A list of extracted phone numbers (core digits only).
        """
        _matches = re.findall(self.PHONE_REGEX, text)
        contact_numbers = []
        for match in _matches:
            match = re.sub(r"\+|\D", "", match)
            
            if len(match) < 10:
                continue
            
            if len(match) > 10:
                match = match[2:]
                
            if match in contact_numbers: 
                continue
            
            contact_numbers.append(match)
            
        contact_numbers = list(set(contact_numbers))
        return contact_numbers

    def extract_emails(self, text):
        """
        Extracts email addresses from text.

        Args:
            text (str): The text to search for email addresses.

        Returns:
            list: A list of extracted email addresses.
        """
        email_matches = re.findall(self.EMAIL_REGEX, text)
        extracted_emails = []
        try:
                
            for email_match in email_matches:
                # Join the two groups to reconstruct the email address
                email = email_match[0].strip() + "@" + email_match[1].strip() + email_match[2].strip()
                extracted_emails.append(email)
                
            extracted_emails = list(set(extracted_emails))
        except Exception as e:
            self.log.exception(f"Unexpected error: {e}", exc_info=True)
        return extracted_emails


    def process_cv(self, cv_path):
        """Extracts email IDs, contact numbers, and overall text from a CV.

        Args:
            cv_path (str): Path to the CV file (PDF or Word).

        Returns:T
            tuple: A tuple containing (email, contact_number, text)
        """


        self.log.info(f"Processing CV: {cv_path}")
        text = ""

        # Determine file type and handle accordingly
        if cv_path.endswith(".pdf"):
            try:
                # Read PDF
                self.log.info(f"Reading PDF: {cv_path}")
                with open(cv_path, 'rb') as f:
                    reader = PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text()
                        
            except (IOError, FileNotFoundError) as e:
                self.log.error(f"Error accessing file: {cv_path} ({e})")
                
            except Exception as e:
                self.log.exception(f"Unexpected error: {e}", exc_info=True)
                    
        elif cv_path.endswith(".docx"):
            try:
                # Read Word document using python-docx
                self.log.info(f"Reading Word: {cv_path}")
                doc = Document(cv_path)  # Use Document() from python-docx

                # Extract text from paragraphs
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text.strip()  # Remove leading/trailing whitespace

                # Handle potential errors more specifically
            except (IOError, FileNotFoundError) as e:
                self.log.error(f"Error accessing file: {cv_path} ({e})")
            except Exception as e:
                self.log.exception(f"Unexpected error: {e}", exc_info=True)

        elif cv_path.endswith(".doc"):
            try:
                # Read DOC using textract (may not be perfect for all DOC formats)
                self.log.info(f"Reading DOC using Spire.doc: {cv_path}")
                doc = SpireDoc(cv_path)
                text = doc.GetText()
                
            except (IOError, FileNotFoundError) as e:
                self.log.error(f"Error accessing file: {cv_path} ({e})")
            except Exception as e:
                self.log.exception(f"Unexpected error: {e}", exc_info=True)

        else:
            self.log.error(f"Unsupported file format: {cv_path}")

        # Extract email IDs and contact numbers using regular expressions (customize patterns as needed)
        email_matches = self.extract_emails(text)

        contact_numbers = self.extract_phone_numbers(text)

        # self.log.info(f"Text: {text}")
        self.log.info(f"Email matches: {email_matches}")
        self.log.info(f"Contact numbers: {contact_numbers}")

        return email_matches, contact_numbers, text


    def find_cvs(self, directory):
        """
        Finds all CV files (PDF and Word documents) in a specified directory and its subdirectories.

        Args:
            directory (str): Path to the directory containing CVs.

        Returns:
            list: A list of filenames (including paths) of all found CVs.
        """

        self.log.info(f"Searching for CVs in {directory}")
        cv_files = []

        try:
            for filename in listdir(directory):
                # Check if it's a directory (avoid hidden directories with leading .)
                if path.isdir(path.join(directory, filename)) and not filename.startswith('.'):
                    # Recursively search for CVs in the subdirectory
                    cv_files.extend(self.find_cvs(path.join(directory, filename)))
                else:
                    # Check for CV file extensions
                    if filename.endswith(".pdf") or filename.endswith(".docx") or filename.endswith(".doc"):
                        cv_path = path.join(directory, filename)
                        self.log.info(f"Found CV file: {cv_path}")
                        cv_files.append(cv_path)

            return cv_files

        except Exception as e:
            self.log.exception(f"Unexpected error {e}", exc_info=True)
            return []  # Return empty list on error

        
    def read_and_exctract_from_cvs(self, directory="CV", zip_file=False):
            
        # Extract ZIP file if necessary
        if zip_file:
            try:
                with zipfile.ZipFile(directory, 'r') as zip_ref:
                    # Extract to a temporary directory to avoid conflicts
                    temp_dir = os.path.join(os.getcwd(), '.temp_extract')
                    os.makedirs(temp_dir, exist_ok=True)  # Create directory if it doesn't exist
                    zip_ref.extractall(temp_dir)
                    directory = temp_dir  # Update directory to temporary location
                    
            except zipfile.BadZipFile:
                self.log.exception(f"Invalid ZIP file: {zip_file}", exc_info=True)
                return {
                'status': 'error',
                'message': 'Invalid ZIP file'
            }
            
            except Exception as e:  # Catch other potential errors
                self.log.exception(f"Error extracting ZIP file: {e}", exc_info=True)
                return {
                'status': 'error',
                'message': f'Error extracting ZIP file {e}'
            }
        
        self.log.info(f"Reading CVs from {directory}")
        cv_files = self.find_cvs(directory)

        # Create a new Excel workbook and worksheet
        workbook = Workbook()
        worksheet = workbook.active

        # Process each CV and write the extracted data to the worksheet
        try:
            worksheet.cell(row=1, column=1).value = "Emails"
            worksheet.cell(row=1, column=2).value = "Contact Numbers"
            worksheet.cell(row=1, column=3).value = "CV content"

            for i, cv_path in enumerate(cv_files, start=1):
                emails, contact_numbers, text = self.process_cv(cv_path)
                
                emails = ','.join(emails)
                contact_numbers = ','.join(contact_numbers)
                
                worksheet.cell(row=i+1, column=1).value = emails
                worksheet.cell(row=i+1, column=2).value = contact_numbers
                worksheet.cell(row=i+1, column=3).value = text

            # Save the workbook
            workbook.save(f"{directory}/output.xlsx")
            
            self.log.info(f"Data is extracted and saved to output.xlsx")
            
            return {
                'status': 'success',
                'excel_file': f"{directory}/output.xlsx"
                }

        except Exception as e:
            self.log.exception(f"Unexpected error {e}", exc_info=True)
            
            return {
                'status': 'error',
                'message': f'Unexpected error {e}'
            }

if __name__ == "__main__":
    scraper = Scraper()
    scraper.read_and_exctract_from_cvs("CV")